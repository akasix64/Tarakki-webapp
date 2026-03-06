"""
Resume / Document Parser API route — uses Gemini AI.
Extracts structured profile data from uploaded PDFs or Word documents.
"""

import os
import json
import tempfile
import requests as http_requests
from flask import Blueprint, request, jsonify
from auth import require_token
import supabase_client as db

resume_parser_bp = Blueprint("resume_parser", __name__)

# ── OpenRouter AI config ────────────────────────────────────────────────────
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "google/gemini-2.0-flash-001"


def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        import PyPDF2
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"[parse-resume] PDF extraction error: {e}")
        return ""


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path)
        parts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Extract from tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    parts.append(" | ".join(row_text))

        # Extract from headers and footers
        for section in doc.sections:
            header = section.header
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())
            footer = section.footer
            if footer:
                for para in footer.paragraphs:
                    if para.text.strip():
                        parts.append(para.text.strip())

        text = "\n".join(parts)
        print(f"[parse-resume] DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, extracted {len(text)} chars")
        return text.strip()
    except Exception as e:
        print(f"[parse-resume] DOCX extraction error: {e}")
        return ""


def _build_prompt(text: str, role: str) -> str:
    """Build the Gemini prompt based on role."""
    if role == "contractor":
        return f"""You are an expert resume parser. Analyze the following resume text and extract structured information.
Return ONLY a valid JSON object with these fields (use null for any field you cannot find):

{{
    "full_name": "Person's full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "website": "LinkedIn URL or personal website",
    "about": "A professional summary (2-3 sentences max)",
    "experience_years": number of years of experience (integer),
    "skills": ["skill1", "skill2", "skill3"]
}}

IMPORTANT RULES:
- Return ONLY the JSON, no markdown formatting, no code blocks, no explanation.
- For experience_years, calculate from the earliest work experience date to present.
- For skills, extract technical skills, tools, technologies, and programming languages as an array.
- For about, write a brief professional summary based on the resume content.
- If a field is not found, set it to null.

RESUME TEXT:
{text}"""
    else:  # startup
        return f"""You are an expert document parser. Analyze the following company document text and extract structured information.
Return ONLY a valid JSON object with these fields (use null for any field you cannot find):

{{
    "full_name": "Company name",
    "email": "Company email address",
    "phone": "Company phone number",
    "location": "Company headquarters location",
    "website": "Company website URL",
    "about": "Company description (2-3 sentences max)",
    "gst_number": "GST Number if found",
    "industry": "Industry/sector the company operates in",
    "company_size": "Company size (one of: 1-10, 11-50, 51-200, 201-500, 500+)",
    "founded_year": founding year (integer)
}}

IMPORTANT RULES:
- Return ONLY the JSON, no markdown formatting, no code blocks, no explanation.
- For company_size, choose the closest option from: 1-10, 11-50, 51-200, 201-500, 500+
- For founded_year, return an integer year.
- If a field is not found, set it to null.

DOCUMENT TEXT:
{text}"""


def _call_ai(prompt: str) -> dict | None:
    """Call OpenRouter AI API and parse the JSON response."""
    api_key = os.getenv("OPENROUTER_API_KEY", "")
    if not api_key:
        print("[parse-resume] ERROR: OPENROUTER_API_KEY is not set!")
        return None

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:3001",
        "X-Title": "Tarakki Resume Parser",
    }

    payload = {
        "model": OPENROUTER_MODEL,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.1,
        "max_tokens": 2048,
    }

    text_response = ""
    try:
        print(f"[parse-resume] Calling OpenRouter AI ({OPENROUTER_MODEL})...")
        resp = http_requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=60)

        if resp.status_code >= 400:
            print(f"[parse-resume] OpenRouter API error: {resp.status_code} — {resp.text[:500]}")
            return None

        data = resp.json()
        # Extract text from OpenAI-compatible response format
        text_response = data["choices"][0]["message"]["content"]

        # Clean up: remove markdown code block markers if present
        text_response = text_response.strip()
        if text_response.startswith("```json"):
            text_response = text_response[7:]
        if text_response.startswith("```"):
            text_response = text_response[3:]
        if text_response.endswith("```"):
            text_response = text_response[:-3]
        text_response = text_response.strip()

        print(f"[parse-resume] AI response parsed successfully")
        return json.loads(text_response)

    except json.JSONDecodeError as e:
        print(f"[parse-resume] Failed to parse AI response as JSON: {e}")
        print(f"[parse-resume] Raw response: {text_response[:500]}")
        return None
    except Exception as e:
        print(f"[parse-resume] OpenRouter API call failed: {e}")
        return None


# ── POST /api/parse-resume ──────────────────────────────────────────────────

@resume_parser_bp.route("/api/parse-resume", methods=["POST"])
def parse_resume():
    """
    Parse a resume/document using Gemini AI.

    Expects JSON body:
    {
        "resume_url": "https://...supabase.co/storage/v1/.../file.pdf",
        "role": "contractor" | "startup"
    }

    Returns parsed profile data as JSON.
    """
    try:
        token = require_token()
        body = request.get_json()

        if not body:
            return jsonify({"error": "Request body is required"}), 400

        resume_url = body.get("resume_url", "").strip()
        role = body.get("role", "contractor").strip().lower()

        print(f"[parse-resume] ── START ──")
        print(f"[parse-resume] Role: {role}")
        print(f"[parse-resume] URL: {resume_url[:120]}...")

        if not resume_url:
            return jsonify({"error": "resume_url is required"}), 400

        # ── Step 1: Download the file ────────────────────────────────────────
        # Include Supabase API key for private/authenticated bucket access
        SUPABASE_KEY = os.getenv("VITE_SUPABASE_ANON_KEY", "")
        download_headers = {}
        if SUPABASE_KEY:
            download_headers["apikey"] = SUPABASE_KEY
            download_headers["Authorization"] = f"Bearer {SUPABASE_KEY}"

        try:
            print(f"[parse-resume] Downloading file...")
            file_resp = http_requests.get(resume_url, headers=download_headers, timeout=30)
            print(f"[parse-resume] Download: status={file_resp.status_code}, content-type={file_resp.headers.get('Content-Type', '?')}, size={len(file_resp.content)} bytes")

            if file_resp.status_code >= 400:
                print(f"[parse-resume] Download FAILED: {file_resp.text[:300]}")
                return jsonify({"error": f"Failed to download file (HTTP {file_resp.status_code})."}), 400
        except Exception as e:
            print(f"[parse-resume] Download exception: {e}")
            return jsonify({"error": f"Failed to download file: {str(e)}"}), 400

        # Determine file type from URL or content type
        content_type = file_resp.headers.get("Content-Type", "")
        # Decode URL-encoded characters for file extension matching
        from urllib.parse import unquote
        decoded_url = unquote(resume_url).lower()

        is_pdf = "pdf" in content_type.lower() or decoded_url.endswith(".pdf")
        is_docx = (
            "wordprocessingml" in content_type.lower()
            or decoded_url.endswith(".docx")
            or decoded_url.endswith(".doc")
        )

        # If we still can't determine type, check content bytes
        if not is_pdf and not is_docx:
            file_bytes = file_resp.content[:4]
            if file_bytes == b'%PDF':
                is_pdf = True
            elif file_bytes == b'PK\x03\x04':
                is_docx = True  # DOCX files are ZIP archives

        print(f"[parse-resume] Detected: is_pdf={is_pdf}, is_docx={is_docx}")

        # ── Step 2: Save to temp file and extract text ───────────────────────
        suffix = ".pdf" if is_pdf else ".docx" if is_docx else ".pdf"

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_resp.content)
            tmp_path = tmp.name

        print(f"[parse-resume] Temp file: {tmp_path}")

        try:
            if is_pdf:
                print(f"[parse-resume] Extracting text from PDF...")
                text = _extract_text_from_pdf(tmp_path)
            elif is_docx:
                print(f"[parse-resume] Extracting text from DOCX...")
                text = _extract_text_from_docx(tmp_path)
            else:
                print(f"[parse-resume] Unknown type — trying both extractors...")
                text = _extract_text_from_pdf(tmp_path)
                if not text:
                    text = _extract_text_from_docx(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        print(f"[parse-resume] Extracted text: {len(text) if text else 0} chars")
        if text:
            print(f"[parse-resume] Preview: {text[:150]}...")
        else:
            print(f"[parse-resume] ERROR: No text extracted!")

        if not text:
            return jsonify({
                "error": "Could not extract text from the document. Make sure it's a valid PDF or Word file with readable text (not scanned images)."
            }), 400

        # Truncate overly long text
        if len(text) > 15000:
            text = text[:15000]

        # ── Step 3: Call Gemini API ──────────────────────────────────────────
        print(f"[parse-resume] Calling Gemini...")
        prompt = _build_prompt(text, role)
        parsed = _call_ai(prompt)

        if not parsed:
            return jsonify({
                "error": "AI could not parse the document. Please fill in the details manually."
            }), 500

        print(f"[parse-resume] Parsed fields: {list(parsed.keys())}")

        # ── Step 4: Clean up and return ─────────────────────────────────────
        if isinstance(parsed.get("skills"), list):
            parsed["skills"] = ", ".join(parsed["skills"])

        if parsed.get("experience_years") is not None:
            parsed["experience_years"] = str(parsed["experience_years"])

        if parsed.get("founded_year") is not None:
            parsed["founded_year"] = str(parsed["founded_year"])

        print(f"[parse-resume] ── SUCCESS ──")
        return jsonify({"parsed": parsed})

    except Exception as e:
        print(f"[parse-resume] ── UNHANDLED ERROR ──")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500
